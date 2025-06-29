import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pygments import highlight
from pygments.lexers import guess_lexer_for_filename
from pygments.styles import get_style_by_name
from pygments.formatter import Formatter

SUPPORTED_EXTENSIONS = [".py", ".sql"]
IGNORED_FILES = {"convert.py", "topdf.py", "toword.py", "screenbot.py"}
OUTPUT_FILE = "project_code.docx"
STYLE_NAME = "friendly"

class CompactDocxFormatter(Formatter):
    def __init__(self, doc, style_name="default", **options):
        super().__init__(**options)
        self.doc = doc
        self.style = get_style_by_name(style_name)
        self.line_num = 1

    def format(self, tokensource, outfile):
        line_tokens = []
        for ttype, value in tokensource:
            lines = value.split('\n')
            for i, part in enumerate(lines):
                line_tokens.append((ttype, part))
                if i < len(lines) - 1:
                    self._write_line(line_tokens)
                    line_tokens = []
        if line_tokens:
            self._write_line(line_tokens)

    def _write_line(self, tokens):
        para = self.doc.add_paragraph()
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.line_spacing = 1.0

        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), "E6E6E6") 
        para._p.get_or_add_pPr().append(shading_elm)

        run = para.add_run(f"{self.line_num:4} | ")
        run.font.name = 'Jetbrains Mono'
        run.font.size = Pt(7)
        run.font.color.rgb = RGBColor(85, 85, 85)

        if not tokens or all(val.strip() == "" for _, val in tokens):
            run = para.add_run(" ")
            run.font.name = 'Jetbrains Mono'
            run.font.size = Pt(7)
        else:
            for ttype, val in tokens:
                run = para.add_run(val)
                run.font.name = 'Jetbrains Mono'
                run.font.size = Pt(7)

                style_def = self.style.style_for_token(ttype)
                color = style_def.get("color")
                run.font.color.rgb = RGBColor.from_string(color) if color else RGBColor(0, 0, 0)

        self.line_num += 1

def get_code_files(root="."):
    files = []
    for dirpath, _, filenames in os.walk(root):
        for f in filenames:
            if f in IGNORED_FILES:
                continue
            if os.path.splitext(f)[1] in SUPPORTED_EXTENSIONS:
                files.append(os.path.join(dirpath, f))
    return files

def generate_docx(files):
    doc = Document()
    doc.add_heading("📘 Project Source Code", 0)

    section = doc.sections[0]
    section.top_margin = Inches(0.3)
    section.bottom_margin = Inches(0.3)
    section.left_margin = Inches(0.4)
    section.right_margin = Inches(0.4)

    for path in files:
        doc.add_page_break()
        doc.add_heading(path, level=2)

        with open(path, 'r', encoding='utf-8') as f:
            code = f.read()

        lexer = guess_lexer_for_filename(path, code)
        formatter = CompactDocxFormatter(doc, style_name=STYLE_NAME)
        highlight(code, lexer, formatter)

    doc.save(OUTPUT_FILE)

if __name__ == "__main__":
    code_files = get_code_files(".")
    if not code_files:
        print("❌ No code files found.")
    else:
        print(f"📂 Found {len(code_files)} code files.")
        generate_docx(code_files)
        print(f"✅ Word file saved as {OUTPUT_FILE}")
